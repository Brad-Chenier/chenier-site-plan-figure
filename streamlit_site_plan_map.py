#!/usr/bin/env python3
"""
Chenier Environmental Consulting
Site Plan Map Generator (Figure 2) — Streamlit Web App

Like the Site Location Map, but:
  - NAIP aerial imagery background (USGS ImageServer exportImage)
  - Always "Figure 2: Site Plan"
  - Portrait OR Landscape (full-page rotation)
  - Same zoom-out slider

Deploy to Streamlit Community Cloud. Files needed in the repo:
  streamlit_site_plan_map.py, requirements.txt, north_arrow.jpeg
"""

import io, math, zipfile
import streamlit as st
import requests
from PIL import Image, ImageDraw
from shapely.geometry import Polygon
from shapely.ops import unary_union
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── NAIP imagery (USGS National Map ImageServer) ────────────────────────────────
NAIP_EXPORT = ('https://imagery.nationalmap.gov/arcgis/rest/services/'
               'USGSNAIPImagery/ImageServer/exportImage')
USGS_HEADERS = {'User-Agent': 'ChenierSitePlan/1.0 (environmental consulting)'}

MIN_BUFFER_DEG = 0.0008      # ~80 m — site plans are tight/close-in
SITE_COLOR     = (255, 0, 0)
JPEG_QUALITY   = 88
NORTH_ARROW_PATH = 'north_arrow.jpeg'

# Layout per orientation. Map image dims chosen to leave room for caption row.
LAYOUTS = {
    'portrait': {
        'orient': WD_ORIENT.PORTRAIT,
        'page_w': Inches(8.5),  'page_h': Inches(11.0),
        'mar_l': Inches(0.7),   'mar_r': Inches(0.25),
        'mar_t': Inches(0.5),   'mar_b': Inches(0.5),
        'map_w': Inches(7.5),   'map_h': Inches(8.5),
        'img_px_w': 1125,       'img_px_h': 1275,
    },
    'landscape': {
        'orient': WD_ORIENT.LANDSCAPE,
        'page_w': Inches(11.0), 'page_h': Inches(8.5),
        'mar_l': Inches(0.5),   'mar_r': Inches(0.5),
        'mar_t': Inches(0.312), 'mar_b': Inches(0.25),
        'map_w': Inches(10.0),  'map_h': Inches(6.4),
        'img_px_w': 1500,       'img_px_h': 960,
    },
}


# ── KMZ / KML ───────────────────────────────────────────────────────────────────
def parse_kmz_bytes(file_bytes, filename):
    if filename.lower().endswith('.kmz'):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            names = [n for n in z.namelist() if n.lower().endswith('.kml')]
            if not names:
                raise ValueError("No KML found inside KMZ")
            kml = z.read('doc.kml' if 'doc.kml' in names else names[0])
    else:
        kml = file_bytes
    root  = etree.fromstring(kml)
    polys = []
    for el in root.iter('{http://www.opengis.net/kml/2.2}coordinates'):
        pts = []
        for tok in el.text.strip().split():
            p = tok.split(',')
            if len(p) >= 2:
                try: pts.append((float(p[0]), float(p[1])))
                except: pass
        if len(pts) >= 3: polys.append(Polygon(pts))
    if not polys: raise ValueError("No polygon found in KMZ/KML")
    g = unary_union(polys)
    return g, g.bounds


# ── Web Mercator helpers (NAIP exportImage works cleanly in 3857) ───────────────
def lnglat_to_merc(lng, lat):
    x = lng * 20037508.34 / 180.0
    y = math.log(math.tan((90 + lat) * math.pi / 360.0)) / (math.pi / 180.0)
    y = y * 20037508.34 / 180.0
    return x, y


def fetch_naip(bounds_merc, px_w, px_h):
    """Request a single NAIP image for the given Web Mercator bbox."""
    xmin, ymin, xmax, ymax = bounds_merc
    params = {
        'bbox': f'{xmin},{ymin},{xmax},{ymax}',
        'bboxSR': 3857, 'imageSR': 3857,
        'size': f'{px_w},{px_h}',
        'format': 'jpgpng', 'f': 'image',
    }
    r = requests.get(NAIP_EXPORT, params=params, headers=USGS_HEADERS, timeout=60)
    r.raise_for_status()
    return Image.open(io.BytesIO(r.content)).convert('RGB')


def draw_boundary_merc(img, site_geom, bounds_merc):
    w, h = img.size
    xmin, ymin, xmax, ymax = bounds_merc
    dx = xmax - xmin
    dy = ymax - ymin

    def to_px(lng, lat):
        mx, my = lnglat_to_merc(lng, lat)
        px = int((mx - xmin) / dx * w)
        py = int((ymax - my) / dy * h)
        return px, py

    draw = ImageDraw.Draw(img)
    lw   = max(3, w // 250)
    geoms = site_geom.geoms if hasattr(site_geom, 'geoms') else [site_geom]
    for poly in geoms:
        if not hasattr(poly, 'exterior'):
            continue
        pts = [to_px(lng, lat) for lng, lat in poly.exterior.coords]
        if len(pts) >= 2:
            draw.line(pts + [pts[0]], fill=SITE_COLOR, width=lw)


# ── Word doc helpers ────────────────────────────────────────────────────────────
def get_or_add(el, tag):
    child = el.find(qn(tag))
    if child is None:
        child = OxmlElement(tag); el.insert(0, child)
    return child

def tbl_border(tbl, val='single', sz=12, color='000000'):
    tblPr = get_or_add(tbl._tbl, 'w:tblPr')
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),val); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)
        b.append(el)
    tblPr.append(b)

def no_border(cell):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto')
        b.append(el)
    tcPr.append(b)

def cell_w(cell, emu):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    el = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(int(emu/914400*1440))); el.set(qn('w:type'),'dxa')
    tcPr.append(el)

def run(para, text, bold=False, italic=False, pt=11, font='Segoe UI', color=None):
    r = para.add_run(text)
    r.bold=bold; r.italic=italic; r.font.name=font; r.font.size=Pt(pt)
    if color: r.font.color.rgb = RGBColor(*color)
    return r


def build_doc_bytes(img_bytes, project_no, orientation, north_arrow_bytes):
    L = LAYOUTS[orientation]
    doc = Document()
    sec = doc.sections[0]
    # Set orientation + dimensions (must set orient AND swap w/h)
    sec.orientation = L['orient']
    sec.page_width  = L['page_w']; sec.page_height = L['page_h']
    sec.left_margin = L['mar_l'];  sec.right_margin = L['mar_r']
    sec.top_margin  = L['mar_t'];  sec.bottom_margin = L['mar_b']
    sec.header_distance = Inches(0); sec.footer_distance = Inches(0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    cw = L['page_w'] - L['mar_l'] - L['mar_r']

    # Map image framed
    mt = doc.add_table(1, 1); tbl_border(mt); mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = get_or_add(mt._tbl, 'w:tblPr')
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(int(cw/914400*1440)))
    tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    mc = mt.cell(0,0); cell_w(mc, cw)
    tcPr = get_or_add(mc._tc,'w:tcPr'); tcMar = OxmlElement('w:tcMar')
    for edge in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:w'),'0'); el.set(qn('w:type'),'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)
    mp = mc.paragraphs[0]; mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0); mp.paragraph_format.space_after = Pt(0)
    mp.add_run().add_picture(io.BytesIO(img_bytes), width=L['map_w'], height=L['map_h'])

    # Caption row: left title/proj, right KEY
    cap_w = cw - Inches(2.25)
    ct = doc.add_table(1, 2); ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = ct.cell(0,0); rc = ct.cell(0,1)
    cell_w(lc, cap_w); cell_w(rc, Inches(2.25))
    no_border(lc); no_border(rc)

    lp1 = lc.paragraphs[0]
    lp1.paragraph_format.space_before = Pt(0); lp1.paragraph_format.space_after = Pt(0)
    run(lp1, 'Figure 2:  Site Plan', bold=True, pt=12)
    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0); lp2.paragraph_format.space_after = Pt(0)
    run(lp2, f'Project No. {project_no}', bold=True, pt=10)

    rp1 = rc.paragraphs[0]
    rp1.paragraph_format.space_before = Pt(0); rp1.paragraph_format.space_after = Pt(0)
    run(rp1, 'KEY:', bold=True, pt=9)
    rp2 = rc.add_paragraph()
    rp2.paragraph_format.space_before = Pt(0); rp2.paragraph_format.space_after = Pt(0)
    if north_arrow_bytes:
        rp2.add_run().add_picture(io.BytesIO(north_arrow_bytes), height=Inches(0.28))
        run(rp2, '  ', pt=9)
    sym = rp2.add_run('━━  '); sym.font.color.rgb = RGBColor(255,0,0); sym.font.size = Pt(12)
    run(rp2, 'Subject Property', pt=9)
    rp3 = rc.add_paragraph()
    rp3.paragraph_format.space_before = Pt(2); rp3.paragraph_format.space_after = Pt(0)
    run(rp3, 'Drawing Not to Scale', italic=True, pt=8)

    # Footer
    ftr = doc.sections[0].footer
    for p in ftr.paragraphs:
        p._element.getparent().remove(p._element)
    ft = ftr.add_table(1, 3, width=cw); ft.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc2 = ft.cell(0,0); mc2 = ft.cell(0,1); rc2 = ft.cell(0,2)
    foot_left = cw - Inches(3.55)
    cell_w(lc2, foot_left); cell_w(mc2, Inches(1.5)); cell_w(rc2, Inches(2.05))
    for c in (lc2,mc2,rc2): no_border(c)
    lf = lc2.paragraphs[0]
    r1 = lf.add_run('Figure 2:  Site Plan'); r1.bold=True; r1.font.name='Segoe UI'
    r1.font.size=Pt(14); r1.font.all_caps=True
    lf2 = lc2.add_paragraph()
    r2 = lf2.add_run(f'Project No. {project_no}'); r2.bold=True
    r2.font.name='Segoe UI'; r2.font.size=Pt(10)
    rf = rc2.paragraphs[0]; rf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf.paragraph_format.space_before = Pt(6)
    cr = rf.add_run('Chenier Environmental Consulting, LLC')
    cr.font.name='Segoe UI'; cr.font.size=Pt(10)
    tP = get_or_add(ft._tbl,'w:tblPr'); brd = OxmlElement('w:tblBorders')
    tp = OxmlElement('w:top'); tp.set(qn('w:val'),'single'); tp.set(qn('w:sz'),'4')
    tp.set(qn('w:color'),'auto'); brd.append(tp); tP.append(brd)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# ── Streamlit UI ─────────────────────────────────────────────────────────────────
st.set_page_config(page_title="Chenier — Site Plan Map", page_icon="🛰", layout="centered")
st.title("🛰 Site Plan Map Generator")
st.caption("Chenier Environmental Consulting, LLC — Figure 2 (NAIP aerial)")
st.markdown("---")

uploaded = st.file_uploader("**1. Upload site boundary (KMZ or KML)**", type=['kmz','kml'])

c1, c2, c3 = st.columns([2, 1.3, 1])
with c1:
    project_no = st.text_input("**2. Project number**", placeholder="e.g. 26-014")
with c2:
    orientation = st.radio("**Orientation**", ['portrait', 'landscape'],
                           format_func=str.title, horizontal=True)
with c3:
    zoom_out = st.slider("**Zoom out**", 0.1, 1.5, 0.4, 0.05,
                         help="Higher = more area around the site")

generate = st.button("⚡ Generate Site Plan Map", type="primary", use_container_width=True)

if generate:
    if not uploaded:
        st.error("Please upload a KMZ or KML file.")
    elif not project_no.strip():
        st.error("Please enter a project number.")
    else:
        try:
            with st.status("Generating site plan...", expanded=True) as status:
                st.write("Reading boundary file...")
                fb = uploaded.read()
                site_geom, bounds = parse_kmz_bytes(fb, uploaded.name)

                L = LAYOUTS[orientation]
                px_w, px_h = L['img_px_w'], L['img_px_h']

                # Buffer the bounds (in degrees), then match the image aspect ratio
                min_lng, min_lat, max_lng, max_lat = bounds
                lng_buf = max((max_lng-min_lng)*zoom_out, MIN_BUFFER_DEG)
                lat_buf = max((max_lat-min_lat)*zoom_out, MIN_BUFFER_DEG)
                b = (min_lng-lng_buf, min_lat-lat_buf, max_lng+lng_buf, max_lat+lat_buf)

                # Convert to Web Mercator, then expand the shorter dimension so
                # the merc bbox aspect ratio matches the target image (no stretch)
                x0,y0 = lnglat_to_merc(b[0], b[1])
                x1,y1 = lnglat_to_merc(b[2], b[3])
                bw, bh = x1-x0, y1-y0
                target_ar = px_w / px_h
                cur_ar = bw / bh
                cx, cy = (x0+x1)/2, (y0+y1)/2
                if cur_ar < target_ar:        # too tall — widen
                    bw = bh * target_ar
                else:                          # too wide — heighten
                    bh = bw / target_ar
                merc = (cx-bw/2, cy-bh/2, cx+bw/2, cy+bh/2)

                st.write("Fetching NAIP aerial imagery...")
                img = fetch_naip(merc, px_w, px_h)

                st.write("Drawing site boundary...")
                draw_boundary_merc(img, site_geom, merc)

                ib = io.BytesIO(); img.save(ib, 'JPEG', quality=JPEG_QUALITY, optimize=True)
                img_bytes = ib.getvalue()

                st.write("Building Word document...")
                na = None
                try:
                    with open(NORTH_ARROW_PATH,'rb') as f: na = f.read()
                except FileNotFoundError:
                    st.write("(north_arrow.jpeg not found — north arrow omitted)")

                safe = project_no.strip().replace('/','_').replace('\\','_')
                docx_bytes = build_doc_bytes(img_bytes, project_no.strip(), orientation, na)
                status.update(label="Done!", state="complete", expanded=False)

            st.success("Site Plan Map generated.")
            st.image(img, caption=f"Preview ({orientation})", use_container_width=True)
            st.download_button(
                "⬇ Download Word Document",
                data=docx_bytes,
                file_name=f"{safe}_Fig_2_Site_Plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary", use_container_width=True)

        except Exception as e:
            st.error(f"Error: {e}")
            import traceback; st.code(traceback.format_exc())

st.markdown("---")
st.caption("Imagery: USDA NAIP via USGS The National Map")
